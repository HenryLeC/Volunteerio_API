from flask import request, jsonify
from API import app, db
from API.database import User, Opportunity, NewUnconfHoursMessages, Logs
from API.auth import token_required
from docx import Document
import pickle
import requests
import datetime
import jwt
import traceback
import io
import json


@app.route("/hours", methods=["POST"])
@token_required
def getHours(user):
    try:
        return jsonify({'hours': str(user.hours)})
    except Exception:
        db.session.add(Logs(traceback.format_exc()))
        db.session.commit()
        return "", 500


@app.route('/addhours', methods=["POST"])
@token_required
def add_hours(user):
    try:
        # Check for hours and reason in request
        try:
            hours = request.form["hours"]
            reason = request.form["reason"]
        except KeyError:
            return jsonify({'msg': 'Hours and reason is required.'}), 500
        id = user.HoursId

        # Increment HoursId by 1
        user.HoursId += 1

        # Add Hours to Unconfirmed List
        UnConfHrs = pickle.loads(user.unconfHours)
        # Add hous and reason to unconfirmed list
        print(pickle.loads(user.unconfHours))
        UnConfHrs.append({
            'id': id,
            'hours': int(hours),
            'reason': reason
        })

        # Add To DB
        user.unconfHours = pickle.dumps(UnConfHrs)
        if len(user.UnconfHoursMessages) == 0:
            Message = NewUnconfHoursMessages()
            user.UnconfHoursMessages.append(Message)
            db.session.add(Message)
        db.session.add(user)
        db.session.commit()

        return jsonify({
            'msg': 'Hours added',
            'unconfHours': pickle.loads(user.unconfHours),
            'confHours': pickle.loads(user.confHours)
        })
    except Exception:
        db.session.add(Logs(traceback.format_exc()))
        db.session.commit()
        return "", 500


@app.route('/Opps', methods=["Post"])
@token_required
def list_opps(user):
    try:
        try:
            dateFilter = request.form["filterDate"]
            nameFilter = "%{}%".format(request.form["filterName"])
        except KeyError:
            return jsonify({'msg': "Please Supply all Paramaters"}), 500

        dateFilter = datetime.datetime.strptime(dateFilter, "%Y-%m-%dT%H:%M:%S")

        Opps = Opportunity.query.join(User).filter(
                                                        User.District == user.District,
                                                        Opportunity.Time > dateFilter,
                                                        Opportunity.Name.like(nameFilter)
                                                    ).order_by(
                                                        Opportunity.Time.desc()
                                                    ).all()
        CleanOpps = []
        for opp in Opps:
            CleanOpps.append({
                "ID": str(opp.id),
                "Name": opp.Name,
                "Location": opp.Location,
                "Hours": opp.Hours,
                "Time": opp.getTime(),
                "Sponsor": User.query.get(int(opp.SponsorID)).name
            })
        return jsonify(CleanOpps)
    except Exception:
        db.session.add(Logs(traceback.format_exc()))
        db.session.commit()
        return "", 500


@app.route('/ClockInOut', methods=["POST"])
@token_required
def Clock(user):
    try:
        try:
            Code = request.form["QrCode"]
        except KeyError:
            return jsonify({
                'msg': 'Please Pass in The Correct Parameters'
            }), 500
        res = False
        RightDict = None
        for Dict in pickle.loads(user.CurrentOpps):
            try:
                if Dict["JWT"] == Code:
                    res = True
                    RightDict = Dict
                    break
            except KeyError:
                pass

        if res:
            Hours = int(round((datetime.datetime.utcnow() -
                        RightDict["StartTime"]).seconds / 3600))
            user.hours += Hours

            OppId = jwt.decode(Code, 'VerySecret', algorithm="HS256")["ID"]
            Opp = Opportunity.query.get(OppId)

            user.PastOpps.append(Opp)
            CurrentOpps = pickle.loads(user.CurrentOpps)
            CurrentOpps.remove(RightDict)
            user.CurrentOpps = pickle.dumps(CurrentOpps)

            db.session.add(user)
            db.session.commit()

            return jsonify({
                'msg': 'Thank You, Your Hours were added.'
            })

        else:
            OppId = jwt.decode(Code, 'VerySecret', algorithm="HS256")["ID"]
            Opp = Opportunity.query.get(OppId)
            if Opp:
                CurrentOpps = pickle.loads(user.CurrentOpps)
                CurrentOpps.append({
                    'StartTime': datetime.datetime.utcnow(),
                    'JWT': Code
                })
                user.CurrentOpps = pickle.dumps(CurrentOpps)

                db.session.add(user)
                db.session.commit()

                return jsonify({
                    'msg': "Thank you for clocking in, don't forget to clock out later"
                })
            else:
                return jsonify({
                    'msg': "Please scan an opportunity code."
                })
    except Exception:
        db.session.add(Logs(traceback.format_exc()))
        db.session.commit()
        return "", 500


@app.route('/BookAnOpp', methods=["POST"])
@token_required
def BookAnOpp(user):
    try:
        try:
            Id = request.form["OppId"]
        except KeyError:
            return jsonify({
                'msg': 'Please Pass in The Correct Parameters'
            }), 500
        Opp = Opportunity.query.get(Id)

        user.BookedOpps.append(Opp)
        db.session.add(user)
        db.session.commit()
        return jsonify({
            'msg': 'Opportunity Booked.'
        })
    except Exception:
        db.session.add(Logs(traceback.format_exc()))
        db.session.commit()
        return "", 500


@app.route('/BookedOpps', methods=["Post"])
@token_required
def BookedOpps(user):
    try:
        Opps = user.BookedOpps
        CleanOpps = []
        for opp in Opps:
            CleanOpps.append({
                "Name": opp.Name,
                "Hours": opp.Hours,
                "Time": opp.Time.strftime("%m/%d/%Y, %H:%M")
            })
        return jsonify(CleanOpps)
    except Exception:
        db.session.add(Logs(traceback.format_exc()))
        db.session.commit()
        return "", 500


@app.route('/PastOpps', methods=["Post"])
@token_required
def PastOpps(user):
    try:
        PastOpps = user.PastOpps
        PastOppsClean = []
        for opp in PastOpps:
            PastOppsClean.append({
                "Name": opp.Name,
                "Hours": opp.Hours,
                "Time": opp.Time.strftime("%m/%d/%Y, %H:%M")
            })
        confHours = pickle.loads(user.confHours)
        HoursClean = []
        for opp in confHours:
            HoursClean.append({
                "Hours": opp["hours"],
                "Reason": opp["reason"],
                "Confirmed": "Confirmed"
            })
        unconfHours = pickle.loads(user.unconfHours)
        for opp in unconfHours:
            HoursClean.append({
                "Hours": opp["hours"],
                "Reason": opp["reason"],
                "Confirmed": "Unconfirmed"
            })

        FullClean = {
            "PastOpps": PastOppsClean,
            "Hours": HoursClean,
        }
        return jsonify(FullClean)
    except Exception:
        db.session.add(Logs(traceback.format_exc()))
        db.session.commit()
        return "", 500


@app.route('/GenerateDoc', methods=["POST"])
@token_required
def GenerateDoc(user: User):
    try:
        email = user.email

        # Helper Function to make cell bold
        def make_row_bold(row):
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
        # GenerateDoc
        opps = user.PastOpps
        ConfHours = pickle.loads(user.confHours)

        Doc = Document("Header.docx")
        Doc.add_heading(f'Total Hours: {user.hours}', level=1)
        table = Doc.add_table(rows=1, cols=4)

        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Name'
        hdr_cells[1].text = 'Date'
        hdr_cells[2].text = 'Hours*'
        hdr_cells[3].text = 'Sponsor'
        make_row_bold(table.rows[0])

        for opp in opps:
            row_cells = table.add_row().cells
            row_cells[0].text = opp.Name
            row_cells[1].text = opp.Time
            row_cells[2].text = opp.Hours
            row_cells[3].text = opp.Sponsor
        for Hour in ConfHours:
            row_cells = table.add_row().cells
            row_cells[0].text = Hour["reason"]
            row_cells[1].text = "NA"
            row_cells[2].text = str(Hour["hours"])
            row_cells[3].text = "NA"

        Doc.add_paragraph("* Hours may be incorrect in Sponsored Opps.")

        # # Testing
        # Doc.save("File.docx")

        # Save To IO Stream
        DocStream = io.BytesIO()
        Doc.save(DocStream)
        DocStream.seek(0)

        # Send Email
        requests.post(
            "https://api.mailgun.net/v3/volunteerio.us/messages",
            auth=("api", json.loads(open("APIKeys.json", "r").read())["MailGun"]),
            files=[("attachment", ("Hours.Docx", DocStream))],
            data={
                "from": "Hours <Hours@volunteerio.us>",
                "to": email,
                "text": "Your Hours",
                "subject": "Your Hours"
            }
        )
        return jsonify({
            'msg': f'Email sent from Hours@Volunteerio.us to {email}.'
        })
    except Exception:
        db.session.add(Logs(traceback.format_exc()))
        db.session.commit()
        return "", 500


@app.route('/Leaderboard', methods=["POST"])
@token_required
def Leaderboard(user):
    try:
        try:
            filt = request.form["filter"]
        except KeyError:
            return jsonify({
                'msg': 'Please Pass in The Correct Parameters'
            }), 500

        if filt == "school":
            users = User.query.filter_by(District=user.District, is_student=True).with_entities(User.name, User.hours).order_by(User.hours.desc()).limit(50).all()
        elif filt == "district":
            users = User.query.filter_by(School=user.School, is_student=True).with_entities(User.name, User.hours).order_by(User.hours.desc()).limit(50).all()

        usersReturn = []
        for i, user in enumerate(users):
            usersReturn.append({
                "name": user.name,
                "hours": user.hours,
                "rank": i+1
            })

        return jsonify(usersReturn)

    except Exception:
        db.session.add(Logs(traceback.format_exc()))
        db.session.commit()
        return "", 500
